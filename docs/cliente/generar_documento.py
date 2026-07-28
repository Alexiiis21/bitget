# -*- coding: utf-8 -*-
"""
Genera el documento para el cliente (Daniel) del Panel de Control Bitget:
  1) Dibuja las ilustraciones con Pillow  -> docs/cliente/img/*.png
  2) Arma el documento con python-docx    -> docs/cliente/PCB-Propuesta-Daniel.docx
  3) Exporta a PDF con Microsoft Word     -> docs/cliente/PCB-Propuesta-Daniel.pdf

Uso:  python docs/cliente/generar_documento.py
"""

import os
from PIL import Image, ImageDraw, ImageFont

# --------------------------------------------------------------------------
# Rutas y paleta
# --------------------------------------------------------------------------

BASE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(BASE, "img")
os.makedirs(IMG, exist_ok=True)

S = 2  # factor de escala para nitidez

INK        = (20, 33, 43)
INK_SOFT   = (86, 100, 113)
INK_FAINT  = (140, 152, 164)
PAPER      = (255, 255, 255)
SURFACE    = (243, 246, 248)
SURFACE_2  = (233, 238, 242)
RULE       = (208, 216, 223)

TEAL       = (14, 124, 107)
TEAL_SOFT  = (223, 240, 236)
STEEL      = (53, 97, 142)
STEEL_SOFT = (226, 235, 244)
AMBER      = (166, 106, 31)
AMBER_SOFT = (250, 240, 224)
GREEN      = (46, 125, 91)
RED        = (174, 68, 54)
VIOLET     = (92, 79, 148)

F = "C:/Windows/Fonts/"


def font(name, size):
    return ImageFont.truetype(F + name, int(size * S))


REG   = "segoeui.ttf"
SEMI  = "seguisb.ttf"
BOLD  = "segoeuib.ttf"
MONO  = "consola.ttf"
MONOB = "consolab.ttf"


# --------------------------------------------------------------------------
# Primitivas de dibujo
# --------------------------------------------------------------------------

def canvas(w, h, bg=PAPER):
    im = Image.new("RGB", (int(w * S), int(h * S)), bg)
    return im, ImageDraw.Draw(im)


def box(d, x, y, w, h, fill=None, outline=None, width=1, radius=6):
    d.rounded_rectangle(
        [x * S, y * S, (x + w) * S, (y + h) * S],
        radius=radius * S, fill=fill, outline=outline,
        width=max(1, int(width * S)),
    )


def txt(d, x, y, s, f, fill=INK, anchor="la", spacing=4):
    d.text((x * S, y * S), s, font=f, fill=fill, anchor=anchor, spacing=spacing * S)


def wrap(d, s, f, max_w):
    words, lines, cur = s.split(), [], ""
    for w in words:
        probe = (cur + " " + w).strip()
        if d.textlength(probe, font=f) <= max_w * S:
            cur = probe
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def para(d, x, y, s, f, max_w, fill=INK_SOFT, lh=17, anchor="la"):
    for i, line in enumerate(wrap(d, s, f, max_w)):
        txt(d, x, y + i * lh, line, f, fill, anchor)
    return y + len(wrap(d, s, f, max_w)) * lh


def arrow(d, x1, y1, x2, y2, color=TEAL, width=2, head=7, dashed=False):
    if dashed:
        import math
        total = math.hypot(x2 - x1, y2 - y1)
        dx, dy = (x2 - x1) / total, (y2 - y1) / total
        pos, on = 0.0, True
        while pos < total - head:
            seg = min(7 if on else 5, total - head - pos)
            if on:
                d.line([(x1 + dx * pos) * S, (y1 + dy * pos) * S,
                        (x1 + dx * (pos + seg)) * S, (y1 + dy * (pos + seg)) * S],
                       fill=color, width=int(width * S))
            pos += seg
            on = not on
    else:
        import math
        total = math.hypot(x2 - x1, y2 - y1)
        dx, dy = (x2 - x1) / total, (y2 - y1) / total
        d.line([x1 * S, y1 * S, (x2 - dx * head) * S, (y2 - dy * head) * S],
               fill=color, width=int(width * S))
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    p = [(x2, y2),
         (x2 - head * math.cos(ang - 0.42), y2 - head * math.sin(ang - 0.42)),
         (x2 - head * math.cos(ang + 0.42), y2 - head * math.sin(ang + 0.42))]
    d.polygon([(px * S, py * S) for px, py in p], fill=color)


def chip(d, x, y, label, f, fg, bg, padx=9, h=22, radius=11):
    w = d.textlength(label, font=f) / S + padx * 2
    box(d, x, y, w, h, fill=bg, radius=radius)
    txt(d, x + w / 2, y + h / 2, label, f, fg, anchor="mm")
    return w


# ==========================================================================
# FIGURA 1 — Cómo se va a ver el panel
# ==========================================================================

def figura_panel():
    W, H = 900, 400
    WIN_H = 332          # alto de la ventana: termina justo despues de la ultima fila
    WIN_BOTTOM = 20 + WIN_H
    im, d = canvas(W, H)

    f_ui     = font(REG, 11)
    f_ui_b   = font(SEMI, 11)
    f_head   = font(SEMI, 9.5)
    f_num    = font(MONO, 11)
    f_numb   = font(MONOB, 11)
    f_grp    = font(BOLD, 10.5)
    f_note   = font(REG, 10)
    f_noteb  = font(SEMI, 10)
    f_title  = font(SEMI, 13)

    # --- ventana de la aplicación ---
    box(d, 24, 20, 620, WIN_H, fill=PAPER, outline=RULE, width=1.2, radius=8)
    box(d, 24, 20, 620, 40, fill=INK, radius=8)
    d.rectangle([24 * S, 48 * S, 644 * S, 60 * S], fill=INK)
    txt(d, 40, 40, "Panel de Control Bitget", f_title, PAPER, anchor="lm")
    txt(d, 628, 40, "● en línea", font(REG, 9.5), (120, 200, 175), anchor="rm")

    # --- barra de acciones ---
    box(d, 24, 60, 620, 42, fill=SURFACE)
    d.line([24 * S, 102 * S, 644 * S, 102 * S], fill=RULE, width=int(1 * S))
    acciones = ["Abrir", "Cerrar", "Take Profit %", "Agregar margen", "Apalancamiento"]
    ax = 40
    for i, a in enumerate(acciones):
        w = d.textlength(a, font=f_ui) / S + 20
        box(d, ax, 71, w, 22, fill=TEAL if i == 2 else PAPER,
            outline=TEAL if i == 2 else RULE, width=1, radius=4)
        txt(d, ax + w / 2, 82, a, f_ui, PAPER if i == 2 else INK, anchor="mm")
        ax += w + 8

    # --- encabezado de la tabla ---
    y = 102
    box(d, 24, y, 620, 34, fill=SURFACE_2)
    d.line([24 * S, (y + 34) * S, 644 * S, (y + 34) * S], fill=RULE, width=int(1.2 * S))
    d.line([236 * S, y * S, 236 * S, (WIN_BOTTOM - 14) * S], fill=RULE, width=int(1 * S))
    d.line([440 * S, y * S, 440 * S, (WIN_BOTTOM - 14) * S], fill=RULE, width=int(1 * S))

    txt(d, 40, y + 17, "SUBCUENTA", f_head, INK_FAINT, anchor="lm")
    txt(d, 338, y + 10, "LONG", f_head, GREEN, anchor="mm")
    txt(d, 542, y + 10, "SHORT", f_head, RED, anchor="mm")
    for cx in (236, 440):
        for j, lab in enumerate(["Liquidación", "Entrada", "Take Profit"]):
            txt(d, cx + 34 + j * 68, y + 25, lab, font(REG, 8.5), INK_FAINT, anchor="mm")

    # --- filas ---
    rows = [
        ("grupo", "CUENTA PRINCIPAL 1", "", ""),
        ("fila", "Sub-01", ("57.240,0", "61.180,5", "63.420,0"), None),
        ("fila", "Sub-02", ("57.190,2", "61.205,0", "63.450,8"), None),
        ("fila", "Sub-07", None, ("68.910,4", "64.720,0", "62.100,5")),
        ("grupo", "CUENTA PRINCIPAL 2", "", ""),
        ("fila", "Sub-03", ("2.410,55", "2.588,10", "2.717,50"), None),
        ("fila", "Sub-11", ("2.409,80", "2.590,00", "2.719,50"),
                            ("2.902,00", "2.744,30", "2.661,00")),
    ]

    y += 34
    for kind, a, b, c in rows:
        if kind == "grupo":
            box(d, 24, y, 620, 26, fill=TEAL_SOFT)
            txt(d, 40, y + 13, a, f_grp, TEAL, anchor="lm")
            y += 26
        else:
            box(d, 24, y, 620, 30, fill=PAPER)
            d.line([24 * S, (y + 30) * S, 644 * S, (y + 30) * S], fill=SURFACE_2, width=int(1 * S))
            txt(d, 48, y + 15, a, f_ui_b, INK, anchor="lm")
            for cx, data, col in ((236, b, GREEN), (440, c, RED)):
                if data is None:
                    txt(d, cx + 102, y + 15, "—", f_num, (205, 212, 219), anchor="mm")
                    continue
                for j, v in enumerate(data):
                    fnt = f_numb if j == 2 else f_num
                    color = col if j == 2 else INK
                    txt(d, cx + 34 + j * 68, y + 15, v, fnt, color, anchor="mm")
            y += 30

    # --- notas laterales ---
    nx = 672
    notas = [
        (118, "Un vistazo, tres datos",
         "Por cada posición solo se muestran precio de liquidación, precio de entrada y "
         "precio de Take Profit. Nada más."),
        (212, "Se actualiza solo",
         "Entre 2 y 5 segundos. Usted no toca nada: los números cambian en pantalla."),
        (292, "Lo que se cierra, desaparece",
         "Cuando una posición se cierra, sus datos se van solos. Si la cuenta queda sin "
         "posiciones, la fila desaparece."),
    ]
    for ny, titulo, cuerpo in notas:
        d.line([(nx - 16) * S, (ny + 8) * S, (nx - 4) * S, (ny + 8) * S], fill=TEAL, width=int(2 * S))
        txt(d, nx, ny, titulo, f_noteb, TEAL)
        para(d, nx, ny + 18, cuerpo, f_note, 200, INK_SOFT, 15)

    im.crop((0, 8 * S, W * S, 372 * S)).save(os.path.join(IMG, "fig1-panel.png"), dpi=(200, 200))


# ==========================================================================
# FIGURA 2 — Cómo viaja la información
# ==========================================================================

def figura_flujo_datos():
    W, H = 900, 400
    im, d = canvas(W, H)

    f_t   = font(SEMI, 13)
    f_st  = font(REG, 10.5)
    f_lab = font(SEMI, 10.5)
    f_b   = font(REG, 10)
    f_tag = font(MONOB, 9)

    # --- dos extremos ---
    box(d, 40, 90, 250, 150, fill=SURFACE, outline=RULE, width=1.2, radius=8)
    txt(d, 165, 122, "Su computadora", f_t, INK, anchor="mm")
    txt(d, 165, 145, "Panel de Control Bitget", f_st, INK_SOFT, anchor="mm")
    box(d, 78, 172, 174, 44, fill=PAPER, outline=TEAL, width=1.4, radius=6)
    txt(d, 165, 187, "Todo ocurre acá:", font(REG, 9.5), INK_FAINT, anchor="mm")
    txt(d, 165, 203, "sus claves nunca salen", f_lab, TEAL, anchor="mm")

    box(d, 610, 90, 250, 150, fill=STEEL_SOFT, outline=STEEL, width=1.2, radius=8)
    txt(d, 735, 122, "Bitget", f_t, INK, anchor="mm")
    txt(d, 735, 145, "Sus cuentas y subcuentas", f_st, INK_SOFT, anchor="mm")
    box(d, 648, 172, 174, 44, fill=PAPER, outline=STEEL, width=1.4, radius=6)
    txt(d, 735, 187, "Es quien manda:", font(REG, 9.5), INK_FAINT, anchor="mm")
    txt(d, 735, 203, "el panel muestra lo que informa", font(REG, 9.5), STEEL, anchor="mm")

    # --- canal 1: en vivo ---
    arrow(d, 605, 135, 300, 135, color=TEAL, width=2.4, head=9)
    chip(d, 372, 96, "CANAL EN VIVO", f_tag, PAPER, TEAL)
    txt(d, 450, 152, "Bitget avisa al instante cuando algo cambia en una cuenta.",
        f_b, INK_SOFT, anchor="mm")
    txt(d, 450, 168, "Como una llamada telefónica que queda abierta todo el día.",
        font(REG, 9.5), INK_FAINT, anchor="mm")

    # --- canal 2: verificación ---
    arrow(d, 300, 205, 605, 205, color=AMBER, width=2.2, head=9, dashed=True)
    chip(d, 358, 218, "VERIFICACIÓN CADA 45 SEGUNDOS", f_tag, AMBER, AMBER_SOFT)
    txt(d, 450, 250, "Por las dudas, el panel vuelve a preguntar todo desde cero.",
        f_b, INK_SOFT, anchor="mm")
    txt(d, 450, 266, "Como pasar lista: si algo no coincide, manda lo que dice Bitget.",
        font(REG, 9.5), INK_FAINT, anchor="mm")

    # --- franja inferior: caída de red ---
    box(d, 40, 300, 820, 72, fill=SURFACE, radius=8)
    d.rectangle([40 * S, 300 * S, 44 * S, 372 * S], fill=AMBER)
    txt(d, 66, 318, "¿Y si se corta internet?", f_lab, INK)
    para(d, 66, 338, "El panel se da cuenta en segundos, se reconecta solo y lo avisa en pantalla. "
                     "Mientras tanto sigue consultando por el segundo camino, así que pierde inmediatez, "
                     "no información. Nunca se queda mostrando datos viejos en silencio.",
         f_b, 760, INK_SOFT, 16)

    im.crop((0, 62 * S, W * S, 384 * S)).save(os.path.join(IMG, "fig2-flujo.png"), dpi=(200, 200))


# ==========================================================================
# FIGURA 3 — Qué pasa cuando usted presiona un botón
# ==========================================================================

def figura_pasos():
    W, H = 900, 400
    im, d = canvas(W, H)

    f_n    = font(MONOB, 16)
    f_t    = font(SEMI, 11)
    f_b    = font(REG, 9.5)
    f_res  = font(SEMI, 11)
    f_big  = font(MONOB, 20)
    f_lab  = font(REG, 9)

    pasos = [
        ("1", "Usted elige", "Una acción, y sobre qué cuentas: una, varias o todas."),
        ("2", "El panel revisa", "Antes de enviar nada verifica que la operación sea posible en cada cuenta."),
        ("3", "Usted confirma", "Ve los precios exactos, cuenta por cuenta, y recién ahí aprueba."),
        ("4", "Se envía en orden", "El panel dosifica los envíos para que Bitget no rechace ninguno."),
        ("5", "Recibe el reporte", "Qué salió bien, qué falló y por qué, cuenta por cuenta."),
    ]

    x, w, gap = 40, 148, 20
    for i, (n, t, b) in enumerate(pasos):
        cx = x + i * (w + gap)
        activo = i == 2
        box(d, cx, 60, w, 138, fill=TEAL_SOFT if activo else SURFACE,
            outline=TEAL if activo else RULE, width=1.4 if activo else 1, radius=8)
        txt(d, cx + 16, 78, n, f_n, TEAL if activo else INK_FAINT)
        txt(d, cx + 16, 108, t, f_t, INK)
        para(d, cx + 16, 130, b, f_b, w - 32, INK_SOFT, 14)
        if i < 4:
            arrow(d, cx + w + 3, 129, cx + w + gap - 3, 129, color=RULE, width=2, head=6)

    # --- tarjeta de resultado ---
    box(d, 40, 232, 820, 132, fill=PAPER, outline=RULE, width=1.2, radius=8)
    box(d, 40, 232, 820, 32, fill=INK, radius=8)
    d.rectangle([40 * S, 252 * S, 860 * S, 264 * S], fill=INK)
    txt(d, 62, 248, "Resultado de la operación   ·   Take Profit 40%   ·   37 cuentas seleccionadas",
        font(SEMI, 10.5), PAPER, anchor="lm")

    métricas = [
        ("31", "correctas", GREEN),
        ("3", "con error", RED),
        ("3", "omitidas", AMBER),
    ]
    mx = 78
    for v, lab, color in métricas:
        txt(d, mx, 296, v, f_big, color, anchor="lm")
        txt(d, mx + d.textlength(v, font=f_big) / S + 10, 299, lab, f_lab, INK_SOFT, anchor="lm")
        mx += 130

    d.line([460 * S, 280 * S, 460 * S, 348 * S], fill=RULE, width=int(1 * S))
    txt(d, 490, 285, "Cada error dice qué pasó, en castellano:", font(SEMI, 9.5), INK)
    txt(d, 490, 303, "Sub-14  ·  margen insuficiente para el monto solicitado", font(MONO, 9), INK_SOFT)
    txt(d, 490, 319, "Sub-22  ·  la cuenta usa margen cruzado, no aplica", font(MONO, 9), INK_SOFT)
    box(d, 490, 336, 190, 22, fill=TEAL, radius=4)
    txt(d, 585, 347, "Reintentar solo las fallidas", font(SEMI, 9.5), PAPER, anchor="mm")

    im.crop((0, 36 * S, W * S, 376 * S)).save(os.path.join(IMG, "fig3-pasos.png"), dpi=(200, 200))


# ==========================================================================
# FIGURA 4 — Protección de credenciales
# ==========================================================================

def figura_seguridad():
    W, H = 900, 390
    im, d = canvas(W, H)

    f_t  = font(SEMI, 13)
    f_b  = font(REG, 11)
    f_c  = font(MONOB, 12.5)
    f_h  = font(SEMI, 14.5)

    # --- capas concéntricas ---
    capas = [
        (40, 60, 340, 260, AMBER_SOFT, AMBER, "3 · Atadas a su Windows",
         "Si copian el archivo a otra computadora, no sirve."),
        (66, 92, 288, 196, (250, 246, 236), AMBER, "2 · Contraseña maestra",
         "Solo usted la conoce. Sin ella, el archivo es ruido."),
        (92, 124, 236, 132, PAPER, TEAL, "1 · Cifrado AES-256",
         "El mismo estándar que usa la banca."),
    ]
    for x, y, w, h, fill, outline, titulo, cuerpo in capas:
        box(d, x, y, w, h, fill=fill, outline=outline, width=1.3, radius=8)
        txt(d, x + 14, y + 10, titulo, font(SEMI, 11), outline)

    box(d, 118, 172, 184, 56, fill=INK, radius=6)
    txt(d, 210, 190, "Sus claves de Bitget", font(SEMI, 11.5), PAPER, anchor="mm")
    txt(d, 210, 210, "API Key · Secret · Passphrase", font(MONO, 9), (150, 170, 185), anchor="mm")

    for x, y, w, h, fill, outline, titulo, cuerpo in capas:
        para(d, x + 14, y + h - 28, cuerpo, font(REG, 9.5), w - 28, INK_SOFT, 13)

    # --- columna derecha ---
    rx = 430
    txt(d, rx, 62, "Tres garantías adicionales", f_h, INK)

    puntos = [
        ("Nunca salen de su computadora",
         "El panel se conecta solo con Bitget. No hay servidor nuestro, no hay nube, "
         "no hay copia en ningún otro lado."),
        ("Nunca se muestran completas",
         "En pantalla siempre aparecen enmascaradas, así:"),
        ("Sin permiso de retiro",
         "Las claves se crean sin la facultad de sacar fondos, y el panel rechaza "
         "cualquiera que sí la tenga. Aunque alguien lograra obtenerlas, no podría "
         "retirar su dinero."),
    ]
    y = 96
    for i, (t, b) in enumerate(puntos):
        d.ellipse([rx * S, (y + 5) * S, (rx + 8) * S, (y + 13) * S], fill=TEAL)
        txt(d, rx + 20, y, t, f_t, INK)
        y = para(d, rx + 20, y + 22, b, f_b, 400, INK_SOFT, 17) + 4
        if i == 1:
            box(d, rx + 20, y, 165, 29, fill=SURFACE, outline=RULE, width=1, radius=4)
            txt(d, rx + 102, y + 15, "BG••••••••4f2a", f_c, INK, anchor="mm")
            y += 38
        else:
            y += 14

    im.crop((0, 42 * S, W * S, 348 * S)).save(os.path.join(IMG, "fig4-seguridad.png"), dpi=(200, 200))


# ==========================================================================
# FIGURA 5 — Cronograma
# ==========================================================================

def figura_cronograma():
    W, H = 900, 410
    im, d = canvas(W, H)

    f_f  = font(MONOB, 10)
    f_t  = font(SEMI, 12)
    f_b  = font(REG, 10)
    f_w  = font(MONO, 10)

    fases = [
        ("FASE 1", "Diseño y arquitectura", "Semanas 1-2",
         "Se define cómo se construye todo. Usted recibe este documento y los planos técnicos.", TEAL, True),
        ("FASE 2", "Conexión con Bitget", "Semanas 3-4",
         "Registro seguro de credenciales y el panel de monitoreo funcionando con datos reales, solo mirando.", STEEL, False),
        ("FASE 3", "Las cinco acciones", "Semanas 5-7",
         "Las cinco operaciones ejecutables, con confirmación previa y reporte por cuenta.", VIOLET, False),
        ("FASE 4", "Pruebas y ajuste fino", "Semana 8",
         "Pruebas con dinero real en montos mínimos, casos de error y ejecutable portable.", AMBER, False),
        ("FASE 5", "Entrega y acompañamiento", "Semana 9",
         "Instalación, manual de uso, capacitación y dos semanas de soporte incluidas.", GREEN, False),
    ]

    x0, y0 = 40, 92
    w, gap, card_h = 152, 12, 200

    d.line([x0 * S, (y0 - 22) * S, (x0 + 5 * (w + gap) - gap) * S, (y0 - 22) * S],
           fill=RULE, width=int(2 * S))

    for i, (f_, titulo, semanas, cuerpo, color, actual) in enumerate(fases):
        cx = x0 + i * (w + gap)
        d.ellipse([(cx + w / 2 - 6) * S, (y0 - 28) * S, (cx + w / 2 + 6) * S, (y0 - 16) * S],
                  fill=color if actual else PAPER, outline=color, width=int(2 * S))

        box(d, cx, y0, w, card_h, fill=SURFACE if not actual else TEAL_SOFT,
            outline=color if actual else RULE, width=1.4 if actual else 1, radius=8)
        d.rectangle([cx * S, y0 * S, (cx + w) * S, (y0 + 4) * S], fill=color)
        txt(d, cx + 14, y0 + 18, f_, f_f, color)
        txt(d, cx + 14, y0 + 38, semanas, f_w, INK_FAINT)
        para(d, cx + 14, y0 + 62, titulo, f_t, w - 28, INK, 18)
        para(d, cx + 14, y0 + 114, cuerpo, f_b, w - 28, INK_SOFT, 15)

    txt(d, x0 + w / 2, y0 + card_h + 18, "▲ estamos acá", font(SEMI, 9.5), TEAL, anchor="mm")

    band_y = y0 + card_h + 40
    box(d, 40, band_y, 820, 50, fill=SURFACE_2, radius=8)
    txt(d, 62, band_y + 15, "Total estimado: 9 semanas de desarrollo + 2 semanas de soporte posterior a la entrega.",
        font(SEMI, 11), INK)
    txt(d, 62, band_y + 33, "Al terminar cada fase usted ve funcionando lo que se construyó, no un informe de avance.",
        font(REG, 10.5), INK_SOFT)

    im.crop((0, 44 * S, W * S, (band_y + 56) * S)).save(
        os.path.join(IMG, "fig5-cronograma.png"), dpi=(200, 200))


# ==========================================================================
# DOCUMENTO
# ==========================================================================

def construir_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- página ---
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for attr, val in (("top_margin", 0.85), ("bottom_margin", 0.85),
                      ("left_margin", 0.95), ("right_margin", 0.95)):
        setattr(sec, attr, Inches(val))

    CONTENT_W = 6.6

    # --- estilos base ---
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.28

    def color(rgb):
        return RGBColor(*rgb)

    def h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI Semibold"
        r.font.size = Pt(16)
        r.font.color.rgb = color(INK)
        rule(p)
        return p

    def rule(p):
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), "0E7C6B")
        borders.append(bottom)
        pPr.append(borders)

    def h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI Semibold"
        r.font.size = Pt(12)
        r.font.color.rgb = color(TEAL)
        return p

    def body(text, size=10.5, italic=False, col=INK):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.italic = italic
        r.font.color.rgb = color(col)
        return p

    def bullet(text, bold_head=None):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if bold_head:
            r = p.add_run(bold_head)
            r.bold = True
            r.font.size = Pt(10.5)
            p.add_run(" ").font.size = Pt(10.5)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        return p

    def figura(nombre, alto_pt=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.add_run().add_picture(os.path.join(IMG, nombre), width=Inches(CONTENT_W))
        return p

    def pie(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(text)
        r.font.size = Pt(8.5)
        r.italic = True
        r.font.color.rgb = color(INK_FAINT)
        return p

    def sombreado(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def destacado(titulo, texto, acento="0E7C6B", fondo="F1F7F5"):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.rows[0].cells[0]
        c.width = Inches(CONTENT_W)
        sombreado(c, fondo)
        c.paragraphs[0].text = ""
        p1 = c.paragraphs[0]
        p1.paragraph_format.space_after = Pt(3)
        r = p1.add_run(titulo)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(acento)
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(texto)
        r2.font.size = Pt(10)
        tcPr = c._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:color"), acento)
        borders.append(left)
        tcPr.append(borders)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def tabla(encabezados, filas, anchos):
        t = doc.add_table(rows=1, cols=len(encabezados))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(encabezados):
            c = t.rows[0].cells[i]
            c.width = Inches(anchos[i])
            sombreado(c, "E9EEF2")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9.5)
        for fila in filas:
            cells = t.add_row().cells
            for i, v in enumerate(fila):
                cells[i].width = Inches(anchos[i])
                p = cells[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                r = p.add_run(v)
                r.font.size = Pt(9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ======================= PORTADA =======================

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("PROPUESTA TÉCNICA · FASE 1")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = color(TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Panel de Control Bitget")
    r.font.name = "Segoe UI Semibold"
    r.font.size = Pt(30)
    r.font.color.rgb = color(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Cómo va a funcionar su sistema para operar\nmúltiples cuentas desde una sola pantalla")
    r.font.size = Pt(14)
    r.font.color.rgb = color(INK_SOFT)

    rule(p)

    for etiqueta, valor in (("Preparado para", "Daniel — propietario del proyecto"),
                            ("Fecha", "Julio de 2026"),
                            ("Documento", "Explicación general, sin lenguaje técnico"),
                            ("Alcance", "Hasta 5 cuentas principales y 100 subcuentas · Futuros USDT-M")):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(etiqueta.upper() + "   ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = color(INK_FAINT)
        r = p.add_run(valor)
        r.font.size = Pt(10.5)
        r.font.color.rgb = color(INK)

    doc.add_paragraph()
    body("Este documento explica, en lenguaje corriente, qué se va a construir, cómo se "
         "verá, cómo se protege su información y en qué plazos. No hace falta ningún "
         "conocimiento técnico para leerlo.", 10.5, True, INK_SOFT)

    doc.add_page_break()

    # ======================= 1 =======================

    h1("Qué vamos a construir")

    body("Hoy, para operar en varias cuentas de Bitget a la vez, hay que entrar a cada una, "
         "buscar la operación, ejecutarla y repetir. Con veinte cuentas eso son veinte "
         "repeticiones de lo mismo; con cien, es directamente inviable en el tiempo que dura "
         "una oportunidad de mercado.")

    body("El Panel de Control Bitget es un programa que se instala en su computadora con "
         "Windows y resuelve exactamente eso: una sola pantalla desde la que se ven todas las "
         "operaciones abiertas de todas las cuentas, y desde la que se ejecuta una acción "
         "sobre una cuenta, sobre un grupo o sobre todas al mismo tiempo.")

    destacado("La idea en una frase",
              "Lo que hoy hace cien veces, lo hace una vez: elige la acción, elige las cuentas, "
              "confirma, y el panel se encarga del resto informándole qué pasó en cada una.")

    h2("Las cinco acciones que va a poder ejecutar")
    for t, b in (("Abrir operaciones.", "Elige el activo, si va en LONG o en SHORT, el monto y el apalancamiento."),
                 ("Cerrar operaciones.", "Cierra lo que está abierto en las cuentas que elija."),
                 ("Poner Take Profit por porcentaje.", "Con botones rápidos para los valores que más usa, o escribiendo el número que quiera."),
                 ("Agregar margen.", "Suma USDT a posiciones ya abiertas para alejar el riesgo de liquidación."),
                 ("Ajustar apalancamiento.", "Cambia el apalancamiento de las cuentas seleccionadas.")):
        bullet(b, t)

    body("Las cinco funcionan igual: se pueden aplicar a una cuenta, a varias o a todas. "
         "Siempre piden confirmación antes de ejecutar, y siempre devuelven un informe de qué "
         "ocurrió en cada cuenta.")

    # ======================= 2 =======================

    h1("Cómo se va a ver")

    body("La pantalla principal es el Centro de Monitoreo: una vista tipo planilla, agrupada "
         "por cuenta principal, donde cada fila es una subcuenta con sus posiciones abiertas. "
         "Está diseñada para leerse de un vistazo, no para estudiarla.")

    figura("fig1-panel.png")
    pie("Figura 1 — El Centro de Monitoreo. Las cifras son ilustrativas.")

    body("La decisión de diseño más importante acá es lo que NO aparece. No hay ganancias ni "
         "pérdidas, no hay porcentajes de rendimiento, no hay gráficos ni indicadores. Solo "
         "los tres números que se necesitan para decidir: a qué precio se liquida la posición, "
         "a qué precio se entró y a qué precio está puesto el Take Profit. Todo lo demás es "
         "ruido que compite por su atención en un momento en que la atención vale dinero.")

    h2("Se mantiene solo")
    for t, b in (("Se actualiza entre 2 y 5 segundos.",
                  "Sin apretar nada, sin recargar."),
                 ("Lo cerrado desaparece.",
                  "Cuando una posición se cierra, sus datos se van solos y sin avisos. Si la cuenta se queda sin posiciones, la fila entera desaparece."),
                 ("Refleja los reposicionamientos.",
                  "Si suma órdenes sobre una posición existente, Bitget fusiona todo y recalcula el precio de entrada y el de liquidación promediados. El panel muestra exactamente lo que Bitget informa: no inventa cuentas propias, y por eso lo que ve en el panel siempre coincide con lo que ve en Bitget.")):
        bullet(b, t)

    # ======================= 3 =======================

    h1("Cómo se comunica su panel con Bitget")

    body("Su panel no guarda ni administra su dinero: Bitget lo hace. Lo que hace el panel es "
         "pedirle información y darle órdenes, usando el canal oficial que Bitget ofrece para "
         "esto. La comunicación va por dos caminos al mismo tiempo, y esa redundancia es "
         "deliberada.")

    figura("fig2-flujo.png")
    pie("Figura 2 — Los dos caminos de comunicación entre el panel y Bitget.")

    body("El primero es un canal en vivo: una conexión que queda abierta y por la que Bitget "
         "avisa en el instante en que algo cambia en una cuenta. Es lo que hace que la pantalla "
         "se sienta inmediata.")

    body("El segundo es una verificación periódica: cada 45 segundos el panel vuelve a "
         "preguntar el estado completo de las cuentas, como quien pasa lista. Sirve para "
         "detectar cualquier aviso que se haya perdido por un problema de red. Si los dos "
         "caminos no coinciden, siempre gana lo que dice Bitget.")

    destacado("Por qué esto le importa",
              "Un panel que dependa de un solo camino puede quedarse mostrando datos viejos sin "
              "que nadie lo note: usted creería que una posición sigue abierta cuando ya se "
              "cerró. Con dos caminos, ese escenario simplemente no ocurre. Es la diferencia "
              "entre una herramienta que se ve bien en una demostración y una en la que se puede "
              "confiar con dinero adentro.",
              acento="35618E", fondo="EEF3F8")

    # ======================= 4 =======================

    h1("Qué pasa cuando presiona un botón")

    body("Ejecutar una orden sobre decenas de cuentas a la vez tiene un riesgo evidente: si "
         "algo está mal, está mal muchas veces. Por eso el recorrido de cada acción tiene "
         "frenos deliberados antes de que salga la primera orden.")

    figura("fig3-pasos.png")
    pie("Figura 3 — El recorrido de una acción, y el informe que recibe al final.")

    h2("Los dos frenos que lo protegen")

    for t, b in (("Revisión previa.",
                  "Antes de enviar nada, el panel comprueba que la operación sea posible en cada cuenta. Si el problema es de la operación en sí —un activo mal escrito, un monto inválido—, no sale ni una sola orden: se lo avisa y no pasa nada."),
                 ("Confirmación con precios reales.",
                  "Cuando pide un Take Profit del 40%, el panel no le pide que confirme «40%»: le muestra el precio exacto que va a quedar puesto en cada cuenta. Usted aprueba precios concretos, no una intención.")):
        bullet(b, t)

    h2("Si algo falla en algunas cuentas")

    body("Puede pasar, y hay que decirlo con claridad: una orden puede entrar en 31 cuentas y "
         "fallar en 3, porque una no tenía margen suficiente o porque Bitget rechazó esa "
         "operación puntual. Cuando ocurre:")

    for t, b in (("Lo correcto queda hecho.",
                  "El panel no deshace las operaciones que salieron bien. Cerrar 31 posiciones válidas porque 3 fallaron sería un daño mayor que el problema original."),
                 ("Se le informa cuenta por cuenta.",
                  "Con el motivo explicado en castellano, no con un código de error."),
                 ("Reintenta solo lo que falló.",
                  "Un botón repite la operación exclusivamente sobre las cuentas con error. El panel está construido de forma tal que un reintento nunca puede duplicar una orden que ya había entrado: es imposible terminar con el doble de posición por reintentar.")):
        bullet(b, t)

    # ======================= 5 =======================

    h1("Cómo se protegen sus credenciales")

    body("Para operar en su nombre, el panel necesita las claves de acceso de cada cuenta. Es "
         "el punto más sensible del proyecto y está tratado como tal.")

    figura("fig4-seguridad.png")
    pie("Figura 4 — Las capas de protección de sus claves.")

    body("Sus claves se guardan cifradas en su propia computadora, protegidas por tres capas "
         "sucesivas: el cifrado AES-256 —el mismo estándar que usa la banca—, una contraseña "
         "maestra que solo usted conoce y que se pide al abrir el programa, y una protección "
         "propia de Windows que ata el archivo a su usuario, de modo que copiarlo a otra "
         "computadora no sirve de nada.")

    destacado("La protección más importante no es técnica",
              "Las claves de acceso se crean sin permiso de retiro, y el panel rechaza "
              "cualquier clave que sí lo tenga. Esto significa que, incluso en el peor "
              "escenario imaginable, quien obtuviera esas claves no podría sacar un solo dólar "
              "de sus cuentas. Es la diferencia entre entregar la llave del auto y entregar la "
              "llave de la guantera.",
              acento="A66A1F", fondo="FAF4EA")

    body("Además: el panel nunca muestra una clave completa en pantalla, ni siquiera para "
         "copiarla; no envía información a ningún servidor nuestro, porque no existe tal "
         "servidor; y solo se comunica con Bitget. No hay estadísticas de uso, ni reportes "
         "automáticos, ni copias en la nube.")

    # ======================= 6 =======================

    h1("Con qué está construido, y por qué le conviene")

    body("Estas son las decisiones técnicas principales, explicadas por lo que significan para "
         "usted y no por su nombre técnico.")

    tabla(
        ["Decisión", "Qué significa para usted"],
        [
            ["Programa de escritorio,\nno una página web",
             "Corre en su computadora, no en internet. No hay un sitio al que alguien pueda entrar, ni una cuenta que puedan robarle. Sus claves nunca salen de su equipo."],
            ["Un solo archivo ejecutable,\nsin instalador",
             "Se copia y se abre. Nada que instalar, nada que configurar. Se puede llevar en un pendrive junto con su configuración."],
            ["Tecnología estándar\ny muy difundida",
             "Está construido con las herramientas más usadas de la industria. Si mañana necesita a otro programador, va a encontrarlo fácilmente: no queda atado a nadie."],
            ["Conexión en vivo\ncon Bitget",
             "La información llega empujada por Bitget en el momento en que cambia, no porque el panel esté preguntando cada tanto. Por eso la pantalla se siente inmediata."],
            ["Envío ordenado\nde las órdenes",
             "Bitget limita cuántas órdenes acepta por segundo. El panel las dosifica automáticamente, así que operar sobre cien cuentas no provoca rechazos ni bloqueos temporales."],
            ["Pruebas automáticas\ndel comportamiento en fallas",
             "Antes de la entrega se simulan cortes de internet, demoras y rechazos de Bitget, para verificar que el panel reacciona bien. Se prueba lo que sale mal, no solo lo que sale bien."],
        ],
        [1.9, 4.7],
    )

    # ======================= 7 =======================

    h1("Qué no va a hacer el panel")

    body("Tan importante como definir qué hace es dejar claro qué queda afuera, para que no "
         "haya sorpresas. Nada de esto es una limitación técnica: es el alcance acordado, y "
         "cualquiera de estos puntos puede incorporarse más adelante como un trabajo aparte.")

    for b in ("No muestra saldos, ganancias, pérdidas ni rendimientos.",
              "No incluye gráficos de precios ni indicadores técnicos.",
              "No guarda un historial consultable con filtros; solo un registro de lo que pasó durante la sesión abierta.",
              "No opera solo ni ejecuta estrategias automáticas: hace exactamente lo que usted le indica, cuando se lo indica.",
              "No funciona en segundo plano ni con el programa cerrado.",
              "No detecta subcuentas automáticamente: usted registra cuáles quiere administrar.",
              "No trabaja con otros exchanges, solo con Bitget.",
              "No funciona en Mac ni en teléfonos: es exclusivamente para Windows."):
        bullet(b)

    # ======================= 8 =======================

    h1("Cronograma")

    figura("fig5-cronograma.png")
    pie("Figura 5 — Las cinco fases del proyecto.")

    body("El criterio de las fases es que en cada una vea algo funcionando. Al terminar la "
         "Fase 2, por ejemplo, va a poder abrir el panel y ver sus posiciones reales en "
         "pantalla, actualizándose solas, aunque todavía no se pueda operar desde ahí. Eso "
         "permite corregir el rumbo temprano, cuando corregir es barato.")

    tabla(
        ["Fase", "Plazo", "Qué recibe al terminarla"],
        [
            ["1 · Diseño y arquitectura", "Semanas 1-2", "Este documento y los planos técnicos completos del sistema."],
            ["2 · Conexión con Bitget", "Semanas 3-4", "El Centro de Monitoreo funcionando con sus cuentas reales, en modo solo lectura."],
            ["3 · Las cinco acciones", "Semanas 5-7", "Todas las operaciones ejecutables, con confirmación e informe por cuenta."],
            ["4 · Pruebas y ajuste", "Semana 8", "Pruebas con dinero real en montos mínimos y el ejecutable portable terminado."],
            ["5 · Entrega", "Semana 9", "Instalación, manual de uso, capacitación y dos semanas de soporte incluidas."],
        ],
        [1.9, 1.1, 3.6],
    )

    # ======================= 9 =======================

    h1("Dos puntos que necesito confirmar con usted")

    body("Durante el diseño de esta fase surgieron dos cuestiones que dependen de cómo están "
         "configuradas hoy sus cuentas en Bitget. Ninguna es un obstáculo, pero conviene "
         "resolverlas ahora y no en el medio del desarrollo.")

    h2("1 · Cada subcuenta necesita sus propias claves")

    body("Bitget no permite que las claves de una cuenta principal operen sobre sus "
         "subcuentas: cada subcuenta es independiente y tiene las suyas. Esto significa que "
         "para administrar cien subcuentas hacen falta cien juegos de claves, no cinco.")

    body("No cambia el funcionamiento del panel, pero sí el trabajo inicial de carga. Para que "
         "no sea una tarea de horas, el panel va a incluir carga masiva: se pegan todas juntas "
         "o se importan desde una planilla, y el panel verifica una por una que funcionen antes "
         "de guardarlas.")

    body("Lo que necesito saber: si sus subcuentas fueron creadas desde la web de Bitget o "
         "mediante herramientas de programación. En el segundo caso existe la posibilidad de "
         "generar las claves automáticamente y evitar la carga manual por completo.")

    h2("2 · El Take Profit por porcentaje necesita una definición suya")

    body("En operaciones con apalancamiento, decir «Take Profit al 50%» es ambiguo, y la "
         "ambigüedad cuesta dinero. Puede significar dos cosas muy distintas:")

    for t, b in (("50% de ganancia sobre lo invertido.",
                  "Con apalancamiento 10x, eso se alcanza cuando el precio se mueve apenas 5%."),
                 ("50% de movimiento del precio.",
                  "El precio tiene que moverse la mitad de su valor: diez veces más lejos que en el caso anterior.")):
        bullet(b, t)

    destacado("Mi recomendación",
              "Usar la primera interpretación —ganancia sobre lo invertido— por defecto, porque "
              "es la que un operador con apalancamiento tiene en la cabeza, y dejar la otra "
              "disponible como opción. En cualquier caso, la pantalla de confirmación siempre "
              "le va a mostrar el precio exacto que va a quedar puesto en cada cuenta, de modo "
              "que la ambigüedad se resuelve mirando, no interpretando.")

    # ======================= 10 =======================

    h1("Qué necesito de su parte")

    for t, b in (("Confirmar los dos puntos anteriores.",
                  "Son los únicos que pueden afectar el plazo si se definen tarde."),
                 ("Una cuenta de prueba en Bitget.",
                  "Con un saldo mínimo, para desarrollar y probar sin tocar sus cuentas reales."),
                 ("Las claves de acceso, sin permiso de retiro.",
                  "Se cargan recién en la Fase 4, cuando el sistema ya está probado."),
                 ("Un rato suyo al final de cada fase.",
                  "Entre veinte y treinta minutos para ver lo construido y decir si va en la dirección correcta.")):
        bullet(b, t)

    doc.add_paragraph()
    p = doc.add_paragraph()
    rule(p)
    body("Quedo a disposición para repasar este documento punto por punto y aclarar cualquier "
         "duda antes de comenzar la Fase 2.", 10, True, INK_SOFT)

    salida = os.path.join(BASE, "PCB-Propuesta-Daniel.docx")
    doc.save(salida)
    return salida


# ==========================================================================
# EXPORTACIÓN A PDF
# ==========================================================================

def exportar_pdf(docx_path):
    import win32com.client as win32
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        doc = word.Documents.Open(docx_path, ReadOnly=False)
        doc.SaveAs2(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(False)
    finally:
        word.Quit()
    return pdf_path


if __name__ == "__main__":
    figura_panel()
    figura_flujo_datos()
    figura_pasos()
    figura_seguridad()
    figura_cronograma()
    print("Figuras generadas.")
    docx = construir_docx()
    print("DOCX:", docx)
    pdf = exportar_pdf(docx)
    print("PDF: ", pdf)
